"""Word (.docx) document parser."""

from __future__ import annotations

import asyncio
from pathlib import Path

from core.models import DocumentType

from .base import BaseParser, ParsedDocument


class WordParser(BaseParser):
    """Parse Word (.docx) documents using python-docx."""

    document_type = DocumentType.WORD
    supported_mimes = [
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        "application/msword",
        "application/vnd.openxmlformats-officedocument.wordprocessingml",
    ]

    async def parse(self, file_path: Path) -> ParsedDocument:
        loop = asyncio.get_running_loop()
        return await loop.run_in_executor(None, self._parse_sync, file_path)

    def _parse_sync(self, file_path: Path) -> ParsedDocument:
        import docx

        doc = docx.Document(str(file_path))

        # Extract paragraphs
        paragraphs = []
        for para in doc.paragraphs:
            if para.text.strip():
                paragraphs.append(para.text)

        # Extract tables
        tables = []
        for i, table in enumerate(doc.tables):
            headers = []
            for cell in table.rows[0].cells:
                headers.append(cell.text)
            tables.append(
                {
                    "index": i,
                    "headers": headers,
                    "rows": len(table.rows) - 1,
                    "text": "\n".join(
                        " | ".join(cell.text for cell in row.cells)
                        for row in table.rows
                    ),
                }
            )

        text = "\n\n".join(paragraphs)
        metadata = {
            "paragraph_count": len(paragraphs),
            "table_count": len(tables),
            "file_type": "docx",
        }

        # Add core properties if available
        if doc.core_properties:
            metadata["author"] = doc.core_properties.author or ""
            metadata["title"] = doc.core_properties.title or ""

        # Also extract images
        images = []
        for rel in doc.part.rels.values():
            if "image" in rel.reltype:
                images.append(Path(rel.target_ref))

        return ParsedDocument(
            text=text,
            metadata=metadata,
            tables=tables,
            images=images,
        )

    def can_handle(self, file_path: Path, mime_type: str) -> bool:
        if mime_type in self.supported_mimes:
            return True
        return file_path.suffix.lower() in (".docx", ".doc")
